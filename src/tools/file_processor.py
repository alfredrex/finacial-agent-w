import os
import re
from typing import List, Optional
from pathlib import Path
import asyncio
from concurrent.futures import ThreadPoolExecutor

from PyPDF2 import PdfReader
from docx import Document
import pandas as pd
from openpyxl import load_workbook

from src.config import settings
from src.state import FileInfo


class FileProcessor:
    def __init__(self):
        self.supported_types = {
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.doc': self._extract_docx,
            '.txt': self._extract_txt,
            '.xlsx': self._extract_xlsx,
            '.xls': self._extract_xlsx,
            '.csv': self._extract_csv,
        }
        self._executor = ThreadPoolExecutor(max_workers=4)
    
    def _get_file_type(self, file_path: str) -> str:
        return Path(file_path).suffix.lower()

    def _is_financial_doc(self, file_path: str, content: str) -> bool:
        """判断是否为金融文档 (PDF/DOCX + 包含财务关键词)。"""
        ext = self._get_file_type(file_path)
        if ext not in ('.pdf', '.docx', '.doc'):
            return False
        keywords = ['营收', '净利润', '证券代码', '董事会', '季度报告', '年度报告',
                    '资产负债表', '利润表', '现金流量', '基本每股收益']
        return any(kw in content for kw in keywords)
    
    def _extract_pdf(self, file_path: str) -> str:
        """提取 PDF 全文（单字符串）。"""
        pages = self._extract_pdf_pages(file_path)
        return "\n".join(text for _, text in pages)

    def _extract_pdf_pages(self, file_path: str) -> list:
        """提取 PDF 逐页内容，返回 [(page_num, text), ...]"""
        try:
            reader = PdfReader(file_path)
            pages = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    pages.append((i + 1, text))
            return pages
        except Exception as e:
            return [(0, f"PDF解析错误: {str(e)}")]
    
    def _extract_docx(self, file_path: str) -> str:
        try:
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    paragraphs.append(row_text)
            return "\n".join(paragraphs)
        except Exception as e:
            return f"DOCX解析错误: {str(e)}"
    
    def _extract_txt(self, file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='gbk') as f:
                return f.read()
        except Exception as e:
            return f"TXT解析错误: {str(e)}"
    
    def _extract_xlsx(self, file_path: str) -> str:
        try:
            wb = load_workbook(file_path, data_only=True)
            all_text = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                all_text.append(f"=== 工作表: {sheet_name} ===")
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    if row_text.strip():
                        all_text.append(row_text)
            return "\n".join(all_text)
        except Exception as e:
            return f"XLSX解析错误: {str(e)}"
    
    def _extract_csv(self, file_path: str) -> str:
        try:
            df = pd.read_csv(file_path, encoding='utf-8')
        except UnicodeDecodeError:
            df = pd.read_csv(file_path, encoding='gbk')
        return df.to_string(index=False)
    
    def extract_text(self, file_path: str) -> str:
        file_type = self._get_file_type(file_path)
        if file_type not in self.supported_types:
            return f"不支持的文件类型: {file_type}"
        return self.supported_types[file_type](file_path)
    
    def chunk_text(self, text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
        """固定大小分块 (兜底方案，非金融文档使用)。"""
        chunk_size = chunk_size or settings.CHUNK_SIZE
        overlap = overlap or settings.CHUNK_OVERLAP
        if len(text) <= chunk_size:
            return [text] if text.strip() else []
        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            if end < len(text):
                last_period = chunk.rfind('。')
                last_newline = chunk.rfind('\n')
                split_point = max(last_period, last_newline)
                if split_point > start + chunk_size // 2:
                    chunk = text[start:split_point + 1]
                    end = split_point + 1
            if chunk.strip():
                chunks.append(chunk.strip())
            start = end - overlap if end < len(text) else end
        return chunks
    
    def chunk_financial(self, text: str) -> List[str]:
        """金融文档 5 步语义分块。"""
        from src.rag.financial_chunker import chunk_financial_text
        return chunk_financial_text(text)

    # ── 财报指标提取 ────────────────────────────────────

    # 表格章节标题检测
    TABLE_HEADERS = [
        "利润表", "损益表", "合并利润表", "合并损益表",
        "资产负债表", "合并资产负债表",
        "现金流量表", "合并现金流量表",
    ]

    # 科目名别名 → 统一名称
    METRIC_ALIASES = {
        "营业收入": "revenue", "营业总收入": "revenue", "主营业务收入": "revenue",
        "营业成本": "cost", "营业总成本": "total_cost",
        "净利润": "net_profit", "归属于母公司所有者的净利润": "net_profit",
        "扣除非经常性损益的净利润": "net_profit_dedup",
        "财务费用": "fin_expense", "销售费用": "sell_expense",
        "管理费用": "admin_expense", "研发费用": "rd_expense",
        "税金及附加": "tax_surcharge", "投资收益": "invest_income",
        "营业利润": "oper_profit", "利润总额": "total_profit",
        "总资产": "total_assets", "资产总计": "total_assets",
        "总负债": "total_liability", "负债合计": "total_liability",
        "所有者权益合计": "net_assets", "股东权益合计": "net_assets",
        "归属于母公司所有者权益合计": "net_assets",
        "货币资金": "cash", "应收账款": "receivables",
        "存货": "inventory", "固定资产": "fixed_assets",
        "短期借款": "short_loan", "长期借款": "long_loan",
        "经营活动产生的现金流量净额": "oper_cf",
        "投资活动产生的现金流量净额": "invest_cf",
        "筹资活动产生的现金流量净额": "fin_cf",
        "基本每股收益": "eps", "稀释每股收益": "eps_diluted",
        "加权平均净资产收益率": "roe",
        "毛利率": "gross_margin", "净利率": "net_margin",
    }

    def _extract_financial_metrics(self, content: str, file_path: str,
                                     pages: list = None) -> tuple:
        """从财报文本提取结构化财务指标 (表格解析模式)。

        优先使用表格行解析 (正则 '科目名 数值1 数值2...' 模式)，
        支持 100+ 指标自动提取，无需逐项硬编码。

        Args:
            content: 全文文本 (用于比率类指标)
            file_path: 文件路径
            pages: 可选 [(page_num, page_text), ...] 用于页码追踪

        Returns:
            (metrics: dict, metrics_pages: dict)
            metrics: {metric_key: raw_value_string}
            metrics_pages: {metric_key: source_page_int}
        """
        metrics = {}
        metrics_pages = {}

        # 如果提供了逐页数据，逐页提取以追踪页码
        if pages:
            source_texts = [(text, num) for num, text in pages]
        else:
            source_texts = [(content, 0)]

        # 1. 表格行模式: "科目名  本期金额  上期金额  变动率"
        table_row = re.compile(
            r'([^\d\n]{2,20}?)\s+'           # 科目名 (2-20个非数字字符)
            r'([\d,\-\.]+\s*[万亿]?元?)'      # 本期金额
        )
        for page_text, page_num in source_texts:
            for m in table_row.finditer(page_text):
                name = m.group(1).strip()
                # 清理科目名中的噪声
                name = re.sub(r'[（(].*?[）)]', '', name).strip()
                name = re.sub(r'其中[：:]', '', name).strip()
                if len(name) < 2 or name in ('项目', '科目', '单位', '元'):
                    continue
                val = m.group(2).strip().replace(',', '').replace(' ', '')
                if val and re.search(r'\d', val):
                    unified = self._normalize_metric_name(name)
                    if unified and unified not in metrics:
                        metrics[unified] = val
                        if page_num > 0:
                            metrics_pages[unified] = page_num

        # 2. 补充: 比率类指标 (百分比格式, 全文匹配)
        ratio_patterns = [
            (r'(?:加权平均)?净资产收益率[：:\s]*([\d.\-]+%?)', 'roe'),
            (r'基本每股收益[：:\s]*([\d.\-]+)', 'eps'),
            (r'稀释每股收益[：:\s]*([\d.\-]+)', 'eps_diluted'),
            (r'毛利率[：:\s]*([\d.\-]+%?)', 'gross_margin'),
            (r'净利率[：:\s]*([\d.\-]+%?)', 'net_margin'),
        ]
        # 有逐页数据时也尝试逐页匹配比率
        if pages:
            for page_text, page_num in source_texts:
                for pattern, key in ratio_patterns:
                    m = re.search(pattern, page_text)
                    if m and key not in metrics:
                        val = m.group(1).strip()
                        metrics[key] = val
                        if page_num > 0:
                            metrics_pages[key] = page_num
        else:
            for pattern, key in ratio_patterns:
                m = re.search(pattern, content)
                if m and key not in metrics:
                    val = m.group(1).strip()
                    metrics[key] = val

        return metrics, metrics_pages

    def _normalize_metric_name(self, name: str) -> str:
        """将中文科目名映射为统一英文 key。"""
        # 精确匹配
        if name in self.METRIC_ALIASES:
            return self.METRIC_ALIASES[name]
        # 模糊匹配: 包含关系
        for alias, key in self.METRIC_ALIASES.items():
            if alias in name or name in alias:
                return key
        # 无法归一化的，用中文名做 key (仍存入)
        return None  # 跳过未知科目，避免噪声
    
    def process_file(self, file_path: str) -> FileInfo:
        """处理文件: 提取文本 → 智能分块 (金融文档走5步语义分块)。"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        content = self.extract_text(file_path)

        # 金融文档 → 语义分块 + 逐页指标提取
        if self._is_financial_doc(file_path, content):
            from src.rag.financial_chunker import chunk_financial_text
            chunks = chunk_financial_text(content)

            # 逐页 PDF 提取 (带页码追踪)
            pages = None
            if self._get_file_type(file_path) == '.pdf':
                try:
                    pages = self._extract_pdf_pages(file_path)
                except Exception:
                    pass
            metrics, metrics_pages = self._extract_financial_metrics(
                content, file_path, pages
            )
        else:
            chunks = self.chunk_text(content)
            metrics = {}
            metrics_pages = {}

        return FileInfo(
            file_path=file_path,
            file_type=self._get_file_type(file_path),
            content=content,
            chunks=chunks,
            metrics=metrics,
            metrics_pages=metrics_pages,
        )
    
    async def process_file_async(self, file_path: str) -> FileInfo:
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self._executor, self.process_file, file_path)
    
    async def process_files_async(self, file_paths: List[str]) -> List[FileInfo]:
        tasks = [self.process_file_async(fp) for fp in file_paths]
        return await asyncio.gather(*tasks)
    
    def process_files(self, file_paths: List[str]) -> List[FileInfo]:
        results = []
        for fp in file_paths:
            try:
                result = self.process_file(fp)
                results.append(result)
            except Exception as e:
                results.append(FileInfo(
                    file_path=fp, file_type="error",
                    content=f"处理错误: {str(e)}", chunks=[]
                ))
        return results


file_processor = FileProcessor()
